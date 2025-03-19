import matplotlib.pyplot as plt
from collections import Counter
from docx import Document
from io import BytesIO
from parser_csv import parser
import json


class Word():
    def __init__(self, path, path_to_word):
        self.path = path
        self.path_to_word = path_to_word
        with open("settings.json", "r", encoding='utf-8') as setts:
            self.setts_json = json.load(setts)
        self.doc = Document()
        self.csv = parser.parse(path)

        for i in list(self.setts_json['Subjects'].keys()):
            self.add_subj(i)
    
    def add_subj(self, subj):
        self.doc.add_heading(subj, 1)
        data1 = self.csv["Предметы"][subj]
        for i in list(data1.keys()):
            if len(self.setts_json["Questions_for_subject"][i]) > 1:
                data = self.csv["Предметы"][subj][i]
                counter = Counter(data)
                total = len(data)
                percentages = [(count / total) * 100 for count in counter.values()]
                labels = counter.keys()
                fig, ax = plt.subplots()
                ax.pie(percentages, labels=labels, autopct='%1.1f%%', startangle=90)
                ax.axis('equal')
                img_stream = BytesIO()
                plt.savefig(img_stream, format='png')
                img_stream.seek(0)
                self.doc.add_heading(i, 2)
                self.doc.add_picture(img_stream)
                img_stream.close()
        
            else:
                data = self.csv["Предметы"][subj][i]
                self.doc.add_heading(i, 2)
                self.doc.add_paragraph("\n".join(map(str, data)))

        for i in self.setts_json["Subjects"][subj]:
            self.add_teacher(i)
        self.doc.save(self.path_to_word)
    

    def add_teacher(self, teach):
        self.doc.add_heading(teach, 1)
        data1 = self.csv["Преподаватели"][teach]
        for i in list(data1.keys()):
            if len(self.setts_json["Questions_for_teachers"][i]) > 1:
                data = self.csv["Преподаватели"][teach][i]
                counter = Counter(data)
                total = len(data)
                percentages = [(count / total) * 100 for count in counter.values()]
                labels = counter.keys()
                fig, ax = plt.subplots()
                ax.pie(percentages, labels=labels, autopct='%1.1f%%', startangle=90)
                ax.axis('equal')
                img_stream = BytesIO()
                plt.savefig(img_stream, format='png')
                img_stream.seek(0)
                self.doc.add_heading(i, 2)
                self.doc.add_picture(img_stream)
                img_stream.close()
            else:
                data = self.csv["Преподаватели"][teach][i]
                self.doc.add_heading(i, 2)
                self.doc.add_paragraph("\n".join(map(str, data)))


# docs = Word("1.csv", "1.docx")

